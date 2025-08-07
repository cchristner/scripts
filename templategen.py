import tkinter as tk
from tkinter import messagebox
from docx import Document
import os
import subprocess
import platform
import tempfile

class WordGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word Document Generator")
        self.root.geometry("400x500")
        
        # Words list - you can modify this with your own words
        self.words = [
            "Introduction", "Background", "Methodology", "Results", 
            "Discussion", "Conclusion", "References", "Appendix",
            "Literature Review", "Objectives", "Hypothesis", "Analysis",
            "Recommendations", "Future Work", "Acknowledgements"
        ]
        
        self.checkboxes = []
        self.var_list = []
        
        # Create UI elements
        self.create_widgets()
    
    def create_widgets(self):
        # Title label
        title_label = tk.Label(self.root, text="Select words to include as headings:", 
                             font=("Arial", 12, "bold"), pady=10)
        title_label.pack()
        
        # Frame for checkboxes with scrollbar
        frame = tk.Frame(self.root)
        frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        canvas = tk.Canvas(frame)
        scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Create checkboxes for each word
        for word in self.words:
            var = tk.IntVar()
            cb = tk.Checkbutton(scrollable_frame, text=word, variable=var, 
                              font=("Arial", 10), pady=2)
            cb.pack(anchor=tk.W)
            self.checkboxes.append(cb)
            self.var_list.append(var)
        
        # Add button
        add_word_btn = tk.Button(self.root, text="Add Custom Word", 
                               command=self.add_custom_word, padx=10)
        add_word_btn.pack(pady=5)
        
        # Generate button
        generate_btn = tk.Button(self.root, text="Generate and Open Word Document", 
                               command=self.generate_and_open_document, 
                               bg="#4CAF50", fg="white", padx=10, pady=5)
        generate_btn.pack(pady=15)
        
    def add_custom_word(self):
        # Create a popup window for adding custom words
        popup = tk.Toplevel(self.root)
        popup.title("Add Custom Word")
        popup.geometry("300x120")
        
        # Center the popup
        popup.geometry("+{}+{}".format(
            self.root.winfo_rootx() + 50,
            self.root.winfo_rooty() + 50))
        
        # Label
        tk.Label(popup, text="Enter a custom word:", pady=10).pack()
        
        # Entry
        entry = tk.Entry(popup, width=25)
        entry.pack(pady=5)
        entry.focus_set()
        
        # Button
        def submit():
            word = entry.get().strip()
            if word:
                self.words.append(word)
                var = tk.IntVar()
                for widget in self.root.winfo_children():
                    if isinstance(widget, tk.Frame):
                        for child in widget.winfo_children():
                            if isinstance(child, tk.Canvas):
                                for frame in child.winfo_children():
                                    cb = tk.Checkbutton(frame, text=word, variable=var, 
                                                     font=("Arial", 10), pady=2)
                                    cb.pack(anchor=tk.W)
                                    self.checkboxes.append(cb)
                                    self.var_list.append(var)
                                    break
                popup.destroy()
            else:
                messagebox.showwarning("Warning", "Please enter a word")
        
        tk.Button(popup, text="Add", command=submit).pack(pady=10)
    
    def generate_and_open_document(self):
        # Get selected words
        selected_words = [self.words[i] for i in range(len(self.words)) 
                        if self.var_list[i].get() == 1]
        
        if not selected_words:
            messagebox.showwarning("Warning", "Please select at least one word")
            return
        
        # Create Word document
        doc = Document()
        
        # Add title
        doc.add_heading("Generated Document", 0)
        
        # Add selected words as headings
        for word in selected_words:
            doc.add_heading(word, level=1)
            doc.add_paragraph("Content for " + word + " goes here.")
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_filename = temp_file.name
            
        # Save to temporary file
        doc.save(temp_filename)
        
        # Open the document with the default application
        try:
            self.open_file(temp_filename)
            messagebox.showinfo("Success", "Document created and opened successfully!")
            
            # Set up deletion of temporary file after closing
            if platform.system() == 'Windows':
                # On Windows, we can't delete the file immediately as it will be in use
                # So we'll let the temp file be cleaned up by the OS eventually
                pass
            else:
                # On Unix systems, we can set up a delayed deletion
                self.root.after(1000, lambda: self.delete_temp_file(temp_filename))
                
        except Exception as e:
            messagebox.showerror("Error", f"Could not open the document: {str(e)}")
            # Clean up the temp file if we couldn't open it
            try:
                os.unlink(temp_filename)
            except:
                pass
    
    def delete_temp_file(self, filepath):
        """Delete the temporary file after a delay to ensure it's opened properly."""
        try:
            os.unlink(filepath)
        except:
            # If file is still in use or already deleted, just ignore
            pass
    
    def open_file(self, filepath):
        """Open a file with the default application based on the operating system."""
        if platform.system() == 'Windows':
            os.startfile(filepath)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', filepath))
        else:  # Linux and other Unix-like
            subprocess.call(('xdg-open', filepath))

if __name__ == "__main__":
    # Create the main window
    root = tk.Tk()
    app = WordGeneratorApp(root)
    root.mainloop()